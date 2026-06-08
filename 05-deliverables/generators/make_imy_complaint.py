# -*- coding: utf-8 -*-
"""Generate a complaint to IMY (Swedish DPA) under Article 77 GDPR for Mr Thaer Saidi.

Writes the .docx into 05-deliverables/ (one level up from this script), so it
works regardless of the current working directory.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Times New Roman"
doc = Document()

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.name = BODY_FONT
    h.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h.font.bold = True
    h.font.size = Pt(12 if i == 1 else 11)

for s in doc.sections:
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)


def p(text="", *, bold=False, italic=False, size=11, align=None,
      space_before=0, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        para.alignment = align
    if text:
        r = para.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY_FONT
    return para


def numbered(num, text, *, indent=0.5):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.first_line_indent = Inches(-indent)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(f"{num}\t"); r.font.name = BODY_FONT; r.font.size = Pt(11)
    r2 = para.add_run(text); r2.font.name = BODY_FONT; r2.font.size = Pt(11)
    return para


def bullet(tag, text, *, indent=1.0, hang=0.5):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.first_line_indent = Inches(-hang)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(f"{tag}\t"); r.bold = True; r.font.name = BODY_FONT; r.font.size = Pt(11)
    r2 = para.add_run(text); r2.font.name = BODY_FONT; r2.font.size = Pt(11)
    return para


def hrule(color="808080"):
    para = doc.add_paragraph()
    pr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pb.append(bottom); pr.append(pb)
    para.paragraph_format.space_after = Pt(4)


def field_table(rows):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(4.3)
    for i, (k, v) in enumerate(rows):
        c0 = tbl.rows[i].cells[0]; c1 = tbl.rows[i].cells[1]
        c0.text = ""; c1.text = ""
        r0 = c0.paragraphs[0].add_run(k); r0.bold = True
        r0.font.name = BODY_FONT; r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(v)
        r1.font.name = BODY_FONT; r1.font.size = Pt(10)
        for cc in (c0, c1):
            cc.paragraphs[0].paragraph_format.space_after = Pt(2)
    return tbl


# ============================ HEADER ============================
p("Integritetsskyddsmyndigheten (IMY)", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
p("Box 8114, 104 20 Stockholm · imy@imy.se", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
hrule()

p("KLAGOMÅL ENLIGT ARTIKEL 77 GDPR", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
p("Complaint to the supervisory authority under Article 77 of the GDPR",
  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

field_table([
    ("Date:", "8 June 2026"),
    ("Complainant’s ref:", "[ ref ]"),
    ("Lodged by:", "[ Name ], Counsel — for and on behalf of the complainant"),
])
p("", space_after=4)

# ============================ 1. COMPLAINANT ============================
doc.add_heading("1.  The complainant (data subject)", level=1)
field_table([
    ("Name:", "Thaer Saidi"),
    ("Personal identity no.:", "[ YYYYMMDD-XXXX ]"),
    ("Address:", "[ home address ]"),
    ("Email:", "[ personal email ]"),
    ("Represented by:", "[ Law firm / advokat ] — see authority enclosed"),
])
p("", space_after=4)

# ============================ 2. CONTROLLERS ============================
doc.add_heading("2.  The organisation(s) complained about", level=1)
numbered("2.1", "This complaint is directed at the following entities, whose precise roles as "
         "controller / joint controller the complainant asks IMY to establish:")
field_table([
    ("Respondent 1:", "Epical Group AB / Epical Sweden AB"),
    ("Org. no.:", "[ org.nr ]"),
    ("Address:", "[ registered address ]"),
    ("Contact:", "Mr Lars Rosenquist (handled the request internally)"),
    ("Respondent 2:", "Nordcloud Hosting Sweden AB"),
    ("Org. no.:", "[ org.nr ]"),
    ("Address:", "[ registered address ]"),
    ("Contact:", "Nordcloud Group Legal / Data Protection Officer"),
])
numbered("2.2", "The extension notice in this matter was issued by Nordcloud Hosting Sweden AB, "
         "while the complainant’s request was made to, and handled internally by, Epical. The "
         "identity of the controller is itself in dispute and forms part of this complaint "
         "(see paragraph 6(a) below).", indent=0.5)

# ============================ 3. SUMMARY ============================
doc.add_heading("3.  Summary of the complaint", level=1)
numbered("3.1", "The complainant is a former member of staff at Epical Sweden AB. On or about "
         "mid-January 2026 he made a request under Article 15 GDPR for access to his personal "
         "data. The controller invoked a two-month extension and undertook to respond by no "
         "later than 27 April 2026. That deadline has passed and no compliant Article 15 "
         "response has been provided.")
numbered("3.2", "In addition to that failure of access, the case material that the complainant has "
         "seen reveals that the organisation has processed, retained and (it appears) disclosed "
         "excessive and highly sensitive personal data of the complainant and of numerous third "
         "parties — including a 105-person spreadsheet of personal identity numbers, salaries, "
         "addresses and trade-union affiliation; BankID identity and signature artifacts of the "
         "complainant and of an unrelated third party; and health/sick-leave information about a "
         "named individual.")
numbered("3.3", "The complainant asks IMY to investigate these matters under Articles 57 and 58 "
         "GDPR, to order the controller to comply with Article 15, and to consider the exercise "
         "of its corrective powers, including an administrative fine under Article 83.")

# ============================ 4. ADMISSIBILITY ============================
doc.add_heading("4.  Jurisdiction and admissibility", level=1)
numbered("4.1", "The complainant is resident in Sweden and the controllers are established in "
         "Sweden. IMY is the competent supervisory authority under Article 55 GDPR. The "
         "complainant is entitled to lodge this complaint under Article 77 GDPR and Chapter 6 of "
         "the Swedish Data Protection Act (lag (2018:218)).")
numbered("4.2", "The complaint is brought within a reasonable time of the infringements becoming "
         "apparent. The complainant has, in parallel, sent a formal letter before action to the "
         "controllers requiring compliance; a copy can be provided to IMY on request.")

# ============================ 5. BACKGROUND ============================
doc.add_heading("5.  Factual background and chronology", level=1)
numbered("5.1", "February 2023 onwards — the period covered by the access request, during which "
         "the complainant’s personal data was processed across Epical, Nordcloud and (per the "
         "extension notice) IBM systems.")
numbered("5.2", "Mid-January 2026 — the complainant submits his Article 15 request, seeking his "
         "complete HR/personnel file, payroll and compensation data, performance records, "
         "internal correspondence (email/Teams/Slack) in which he is the subject or is "
         "referenced, system/access/audit logs, and the Article 15(1)(a)–(h) information.")
numbered("5.3", "13 February 2026 — Nordcloud Hosting Sweden AB issues an extension notice. It "
         "states the request was “submitted on 27.1.2026”, asserts the case is complex, and "
         "undertakes to respond “by no later than 27.4.2026”. The internal Epical dossier, by "
         "contrast, records receipt in “mid-January” and the matter as handled by Mr Lars "
         "Rosenquist.")
numbered("5.4", "27 April 2026 — the controller’s own extended deadline. It passed without a "
         "compliant Article 15 response.")
numbered("5.5", "The internal DSAR dossier records the data-source inventory as still “Pending "
         "IT/HR input” (HR systems, email systems, Teams/collaboration and IT logs all marked "
         "“Pending”), and asserts blanket exclusions of entire mailboxes, full email/Teams "
         "threads and attachments.")

# ============================ 6. INFRINGEMENTS ============================
doc.add_heading("6.  The infringements complained of", level=1)
p("The complainant relies on the following infringements, each cross-referenced to the exhibits "
  "in the Schedule.", space_after=8)

bullet("(a)", "Controller identity / accountability (Arts 5(2), 12, 13–14). The organisation has "
              "not given a clear, consistent account of which entity is the controller, "
              "undermining the complainant’s ability to exercise his rights.")
bullet("(b)", "Failure to comply with the right of access within the time limit (Arts 12(3), "
              "15(1), 15(3)). Even on the controller’s own extended deadline of 27 April 2026, no "
              "compliant response was provided; the data-source inventory remained “Pending”. "
              "[Exhibits D, E]")
bullet("(c)", "Improper reliance on blanket exclusions without reasons (Arts 12(4), 15). The "
              "controller asserted categorical refusals of mailboxes, threads and attachments "
              "instead of extracting/summarising the complainant’s personal data and giving "
              "specific reasons for any redactions. [Exhibit E]")
bullet("(d)", "Unlawful / excessive processing and apparent disclosure of third-party personal "
              "data (Arts 5(1)(c), 5(1)(f), 6, 9, 32). A single spreadsheet exposes 105 "
              "employees’ personal identity numbers, private contact details, home addresses, "
              "salaries and trade-union club/section — trade-union membership being "
              "special-category data under Article 9. [Exhibit A]")
bullet("(e)", "Exposure of BankID identity and signature artifacts (Arts 5(1)(c), 5(1)(f), 32; "
              "Arts 33–34). The material contains BankID completion records for the complainant "
              "and an unrelated third party, including personal identity numbers, device/IP "
              "identifiers and full signature/certificate payloads. [Exhibit B]")
bullet("(f)", "Excessive sensitive HR/financial/family data (Arts 5(1)(c), 9, 15(1)(b), 32). An "
              "export of the complainant’s Unit4 record includes his IBAN/bank account, absence "
              "records and next-of-kin (spouse) details. [Exhibit C]")
bullet("(g)", "Health data of third parties in broad exports (Arts 9, 5(1)(f)). An export "
              "discloses a named individual’s heart attack and sick leave. [Exhibit F]")
bullet("(h)", "Inadequate security and data-protection-by-design (Arts 25, 32). The corpus mixes "
              "HR/identity data with operational credentials (forwarded SFTP credentials; Azure "
              "role assignments), evidencing inadequate technical and organisational measures. "
              "[Exhibit G]")
bullet("(i)", "Possible unnotified personal data breach (Arts 33, 34). If the third-party data at "
              "(d)–(g) was disclosed to the complainant or was broadly accessible internally, the "
              "controller’s breach-notification duties are engaged and appear not to have been "
              "met.")

# ============================ 7. RELIEF ============================
doc.add_heading("7.  What the complainant asks IMY to do", level=1)
numbered("7.1", "Exercise its investigative powers under Article 58(1) to obtain from the "
         "controllers a full account of the processing, the data-source inventory, the exact "
         "receipt date of the request, the controller’s identity, and the lawful basis and "
         "safeguards applied to the third-party special-category, financial and identity data.")
numbered("7.2", "Issue an order under Article 58(2)(c) requiring the controller to comply with the "
         "complainant’s Article 15 request, including by providing a copy of his personal data "
         "(with intelligible extracts/summaries of relevant messages and logs) and all "
         "Article 15(1)(a)–(h) information, and to give Article 12(4) reasons for any redactions.")
numbered("7.3", "Investigate whether the member/union list, the BankID artifacts and the broad "
         "exports constitute a personal data breach, and whether the controller complied with "
         "Articles 33 and 34.")
numbered("7.4", "Consider the exercise of its corrective powers under Article 58(2), including a "
         "reprimand, a compliance order and, if appropriate, an administrative fine under "
         "Article 83, having regard to the sensitivity and volume of the data and the number of "
         "data subjects affected.")
numbered("7.5", "Keep the complainant informed of the progress and outcome of the complaint, as "
         "required by Article 77(2) GDPR.")

# ============================ 8. EVIDENCE ============================
doc.add_heading("8.  Evidence relied upon (Schedule of exhibits)", level=1)
p("The following exhibits are available to be produced to IMY. Third-party identifiers are "
  "redacted in the complainant’s working copies and unredacted originals can be supplied to IMY "
  "on a confidential basis.", space_after=8)
exhibits = [
    ("Exhibit A", "Spreadsheet of 105 employees — personnummer, contact details, addresses, "
                  "salaries, trade-union club/section (“member-list”)."),
    ("Exhibit B", "BankID signature/SDO records for the complainant and a third party — personal "
                  "identity numbers, device/IP identifiers, signature/certificate payloads."),
    ("Exhibit C", "Unit4 export of the complainant’s HR record — personnummer, bank account/IBAN, "
                  "absence records, next-of-kin (spouse) details."),
    ("Exhibit D", "Extension notice dated 13 February 2026 (Nordcloud) — receipt 27.1.2026, "
                  "deadline 27.4.2026."),
    ("Exhibit E", "Internal Epical DSAR dossier — “Pending” inventory, “mid-January” receipt, "
                  "handling by Mr Lars Rosenquist, blanket scope exclusions."),
    ("Exhibit F", "Teams/email export disclosing a named third party’s heart attack / sick leave."),
    ("Exhibit G", "Teams export forwarding SFTP credentials; Azure role-assignment export."),
]
for tag, text in exhibits:
    bullet(tag + ":", text, indent=1.3, hang=1.3)

# ============================ 9. DECLARATION ============================
doc.add_heading("9.  Declaration", level=1)
numbered("9.1", "The complainant confirms that the information given in this complaint is true to "
         "the best of his knowledge and belief, and consents to IMY processing the personal data "
         "in this complaint for the purpose of investigating it.")
p("", space_after=18)
p("_____________________________", space_after=0)
p("Thaer Saidi (complainant)", bold=True, space_after=10)
p("_____________________________", space_after=0)
p("[ Name ], [ Advokat / title ]", bold=True, space_after=0)
p("for and on behalf of [ LAW FIRM NAME ] — Counsel for the complainant", italic=True, space_after=14)

# ============================ ATTACHMENTS / NOTE ============================
doc.add_heading("Attachments", level=1)
for a in [
    "1.  Letter before action to Epical / Nordcloud dated 8 June 2026.",
    "2.  Power of attorney / authority to act.",
    "3.  Exhibits A–G (as listed in section 8).",
    "4.  Copy of the complainant’s identification.",
]:
    bullet("", a, indent=0.5, hang=0.0)

p("", space_after=8)
hrule()
p("DRAFTING NOTE (delete before filing): Complete all [ ] fields (complainant ID/address, "
  "org. numbers, addresses, firm details, signatory). IMY normally conducts proceedings in "
  "Swedish — consider filing in Swedish or enclosing a Swedish translation. Verify the true "
  "DSAR receipt date and the controller entity, and confirm what (if anything) has already been "
  "disclosed to the complainant, before filing. Prepared on the instructions and material "
  "provided; to be reviewed by the responsible advokat. This is not, by itself, legal advice.",
  italic=True, size=9)

out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..",
                                   "Complaint_to_IMY_Thaer_Saidi.docx"))
doc.save(out)
print("Saved:", out)
print("Paragraphs:", len(doc.paragraphs), "| Tables:", len(doc.tables))
