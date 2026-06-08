# -*- coding: utf-8 -*-
"""Generate a Letter Before Action (.docx) for Mr Thaer Saidi (GDPR / DSAR).

Writes the .docx into 05-deliverables/ (one level up from this script), so it
works regardless of the current working directory.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Times New Roman"

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.name = BODY_FONT
    h.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h.font.bold = True
    h.font.size = Pt(12 if i == 1 else 11)

# margins
for s in doc.sections:
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)


def p(text="", *, bold=False, italic=False, size=11, align=None,
      space_before=0, space_after=6, font=BODY_FONT):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        para.alignment = align
    if text:
        r = para.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = font
    return para


def runs(para_runs, *, align=None, space_after=6, space_before=0):
    """Add a paragraph built from a list of (text, kwargs) run tuples."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        para.alignment = align
    for text, kw in para_runs:
        r = para.add_run(text)
        r.bold = kw.get("bold", False)
        r.italic = kw.get("italic", False)
        r.font.size = Pt(kw.get("size", 11))
        r.font.name = BODY_FONT
    return para


def numbered(num, text, *, lead_bold=None):
    """Manual legal numbering, e.g. '1.1' with a hanging indent."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.5)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(f"{num}\t")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    if lead_bold:
        rb = para.add_run(lead_bold)
        rb.bold = True
        rb.font.name = BODY_FONT
        rb.font.size = Pt(11)
    r2 = para.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(11)
    return para


def hrule():
    para = doc.add_paragraph()
    pr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pb.append(bottom)
    pr.append(pb)
    para.paragraph_format.space_after = Pt(4)


# ============================ LETTERHEAD ============================
p("[ LAW FIRM NAME ]", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
p("Advokat / Data Protection & Information Law", italic=True, size=9,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
p("[ Street address ] · [ Postcode, City ] · Sweden · [ email ] · [ telephone ]",
  size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
hrule()

p("STRICTLY PRIVATE & CONFIDENTIAL — ADDRESSEE ONLY", bold=True, size=9, space_after=0)
p("BY EMAIL AND BY RECORDED POST", bold=True, size=9, space_after=10)

# ---- addressee / sender block as a 2-col table ----
tbl = doc.add_table(rows=1, cols=2)
tbl.autofit = True
left, right = tbl.rows[0].cells

def cell_lines(cell, lines):
    cell.text = ""
    first = True
    for line, bold in lines:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        rr = para.add_run(line)
        rr.bold = bold
        rr.font.name = BODY_FONT
        rr.font.size = Pt(10)
        first = False

cell_lines(left, [
    ("To:", True),
    ("Epical Group AB", True),
    ("(and the relevant group company, Epical Sweden AB)", False),
    ("Attn: Data Protection Officer / Head of Legal", False),
    ("[ registered address ]", False),
    ("", False),
    ("And to:", True),
    ("Nordcloud Hosting Sweden AB", True),
    ("Attn: Nordcloud Group Legal / Data Protection Officer", False),
    ("[ registered address ]", False),
])
cell_lines(right, [
    ("Date:", True),
    ("8 June 2026", False),
    ("", False),
    ("Our ref:", True),
    ("[ matter / file ref ]", False),
    ("", False),
    ("Your ref:", True),
    ("DSAR – Thaer Saidi;", False),
    ("Extension notice dated 13 February 2026", False),
])

p("", space_after=4)
hrule()

# ============================ RE LINE ============================
runs([
    ("Dear Sirs / Data Protection Officer,", {}),
], space_after=10)

runs([
    ("RE:  Mr Thaer Saidi — Subject Access Request under Article 15 of the GDPR; "
     "failure to comply within the statutory period; unlawful and excessive processing "
     "and disclosure of personal data; and ", {"bold": True}),
    ("NOTICE BEFORE COMPLAINT AND LEGAL PROCEEDINGS", {"bold": True}),
], space_after=10)

# ============================ 1. INTRODUCTION ============================
doc.add_heading("1.  Introduction and capacity", level=1)
numbered("1.1", "We act for Mr Thaer Saidi (“our client”), a former member of staff at Epical "
         "Sweden AB whose personal data is, and has been, processed by your organisation(s). All "
         "correspondence in this matter should now be directed to this firm and to our client’s "
         "nominated personal email address, his access to internal systems and to his "
         "@epicalgroup.com mailbox having been revoked.")
numbered("1.2", "This is a formal letter before action. Its purpose is to put you on notice of "
         "our client’s claims arising under the EU General Data Protection Regulation 2016/679 "
         "(“GDPR”) and the Swedish Data Protection Act (lag (2018:218)), to require you to "
         "remedy the matters set out below, and to give you a final opportunity to do so before "
         "our client lodges a complaint with the Integritetsskyddsmyndigheten (“IMY”) and "
         "commences proceedings for compensation under Article 82 GDPR.")
numbered("1.3", "This is an open letter. We reserve the right to place it, and any response or "
         "failure to respond, before IMY and/or a competent court on the question of liability "
         "and costs.")

# ============================ 2. CONTROLLER / PARTIES ============================
doc.add_heading("2.  The controller(s) and the unresolved identity question", level=1)
numbered("2.1", "Your own correspondence is internally inconsistent as to who the data controller "
         "is. Our client submitted his access request to Epical Group, and your internal records "
         "show the request being “forwarded internally and handled by” Mr Lars Rosenquist of "
         "Epical. The statutory extension notice, however, was issued in the name of “Nordcloud "
         "Hosting Sweden AB / Nordcloud Group Legal / DPO” and asserts that the data is "
         "processed by Nordcloud across “Nordcloud and IBM systems”.")
numbered("2.2", "Article 13/14 and the accountability principle in Article 5(2) require the "
         "controller’s identity to be clear. Our client is entitled to know which legal entity "
         "is the controller of his personal data, and which entity is answerable for the failures "
         "described below. We therefore address this letter to both Epical Group AB / Epical "
         "Sweden AB and Nordcloud Hosting Sweden AB, and we require each of you to confirm, in "
         "writing, your respective roles as controller, joint controller or processor in relation "
         "to our client’s personal data.")

# ============================ 3. THE REQUEST AND THE DEADLINE ============================
doc.add_heading("3.  The access request and the expired statutory deadline", level=1)
numbered("3.1", "Our client made a request under Article 15 GDPR for access to his personal data, "
         "covering the period from February 2023. The request expressly sought, among other "
         "things: his complete personnel and HR file; performance and appraisal records; payroll "
         "and compensation information; internal correspondence and records (including email, "
         "Slack/Teams messages and meeting notes) in which he is the subject or is specifically "
         "referenced; and all system, access, identity and audit logs associated with his user "
         "identity. He further requested the information required by Article 15(1)(a)–(h): the "
         "purposes of processing, the categories of personal data, the recipients, the retention "
         "periods, the source of the data, and the safeguards applied.")
numbered("3.2", "Your records date receipt of the request to “mid-January 2026”, while your "
         "extension notice states it was “submitted on 27.1.2026”. By the notice dated 13 "
         "February 2026 you purported to invoke the two-month extension under Article 12(3) GDPR, "
         "undertaking to respond “by no later than 27.4.2026”.")
numbered("3.3", "Two points arise. First, the discrepancy between “mid-January” and “27 "
         "January” is material: Article 12(3) requires the controller to inform the data subject "
         "of any extension within one month of receipt, and the lawfulness and timeliness of your "
         "extension turns on the true receipt date, which you must now confirm with evidence. "
         "Second, and in any event, ", lead_bold=None)
numbered("3.4", "even taking your own extended deadline of 27 April 2026 at its highest, that "
         "deadline has now passed. As at the date of this letter our client has not received a "
         "compliant Article 15 response. You are therefore in breach of Articles 12(3) and 15 "
         "GDPR.", lead_bold=None)

# ============================ 4. BREACHES ============================
doc.add_heading("4.  The breaches identified", level=1)
numbered("4.1", "From the material available to our client, and from your own internal dossier, the "
         "following breaches are apparent. We summarise them here; the supporting exhibits are "
         "listed in the Schedule to this letter.")

doc.add_heading("(a)  Incomplete and pre-limited Article 15 response", level=2)
numbered("4.2", "Your internal DSAR dossier records the data-source inventory as still “Pending "
         "IT/HR input” — “HR systems: Pending”, “Email systems: Pending”, "
         "“Teams/Collaboration: Pending”, “IT logs: Pending” — at or after the point the "
         "statutory period was due to expire. It also asserts blanket exclusions of “entire "
         "email mailboxes or raw system exports”, “full email conversations or Teams threads” "
         "and “attachments or documents”. While a controller may redact genuine third-party "
         "data and business secrets, it must still identify and provide the data subject’s own "
         "personal data (including extracts or summaries of relevant messages) and the Article "
         "15(1)(a)–(h) information, and must give reasons under Article 12(4) for anything it "
         "withholds. A pending inventory and categorical refusals are not an Article 15 response. "
         "(Articles 12(3), 12(4), 15(1), 15(3), 5(2).)")

doc.add_heading("(b)  Unlawful and excessive disclosure of third-party personal data", level=2)
numbered("4.3", "The case material and exports include a single spreadsheet listing 105 identified "
         "third-party employees together with their Swedish personal identity numbers "
         "(personnummer), private email addresses, mobile numbers, home addresses, monthly "
         "salaries, employment dates and trade-union club/section affiliation. Trade-union "
         "membership is special-category data under Article 9 GDPR. The presence and handling of "
         "this list — whether disclosed to our client or held in a broadly accessible "
         "SharePoint/Teams context — is a serious failure of data minimisation and "
         "confidentiality. (Articles 5(1)(c), 5(1)(e), 5(1)(f), 6, 9, 32.)")

doc.add_heading("(c)  Exposure of BankID identity and signature artifacts", level=2)
numbered("4.4", "The material includes BankID “SDO” completion records for our client and for at "
         "least one other, unrelated individual. Each record contains the signer’s name and "
         "personal identity number, device identifier (uhi), IP address, order reference and the "
         "full XML signature/certificate payload. Including another person’s BankID identity and "
         "signature material in our client’s case file, and handling our client’s own such "
         "material in this way, is a grave confidentiality and security failure. (Articles "
         "5(1)(c), 5(1)(f), 6, 32; and see Articles 33–34.)")

doc.add_heading("(d)  Excessive sensitive HR, payroll, financial and family data", level=2)
numbered("4.5", "An export of our client’s Unit4 record contains his personal identity number, "
         "bank account and IBAN, employment terms, absence records (including leave data), and "
         "next-of-kin details naming his wife together with her address and mobile number. To the "
         "extent this is provided as a disclosure, it must be accompanied by the full Article 15 "
         "metadata; to the extent it sits in an unprotected working folder, it engages data "
         "minimisation and Article 32 security. (Articles 5(1)(c), 5(1)(f), 9, 15(1)(b), 32.)")

doc.add_heading("(e)  Health/sick-leave data of third parties in broad exports", level=2)
numbered("4.6", "The exports include a Teams/email message disclosing that a named individual “had "
         "a minor heart attack last week so he will be on sick leave”. Health data attracts the "
         "heightened protection of Article 9. Its presence in a broad, multi-recipient export "
         "without strict necessity or redaction is a further Article 9 and Article 5(1)(f) "
         "concern.")

doc.add_heading("(f)  Insecure handling and mixing of credentials with personal data", level=2)
numbered("4.7", "The corpus mixes HR and identity material with operational security information, "
         "including a message forwarding “the SFTP credentials for basware” and an export of "
         "Azure role assignments. This demonstrates an absence of appropriate technical and "
         "organisational measures and data-protection-by-design, contrary to Articles 25 and 32, "
         "and illustrates why unreviewed raw exports should never have been generated.")

numbered("4.8", "Taken together, the above indicate breaches of GDPR Articles 5(1)(c), 5(1)(f), "
         "5(2), 6, 9, 12, 15, 25 and 32, and they raise a personal-data-breach assessment duty "
         "under Articles 33 and 34. We put you on notice that, if the third-party data described "
         "above was disclosed to our client or was broadly accessible internally, you may already "
         "be in breach of your Article 33/34 obligations.", lead_bold=None)

# ============================ 5. WHAT WE REQUIRE ============================
doc.add_heading("5.  What we require you to do", level=1)
numbered("5.1", "Within ", lead_bold=None)
# emphasise the deadline inline
dl = doc.paragraphs[-1]
rr = dl.add_run("14 days")
rr.bold = True
rr.font.name = BODY_FONT
rr.font.size = Pt(11)
rr2 = dl.add_run(" of the date of this letter, we require each addressee to:")
rr2.font.name = BODY_FONT
rr2.font.size = Pt(11)

reqs = [
    ("(a)", "confirm in writing which legal entity is the controller of our client’s personal "
            "data, and the respective roles of Epical and Nordcloud (controller, joint controller "
            "or processor);"),
    ("(b)", "confirm the true date on which the access request was received and provide evidence "
            "that any Article 12(3) extension was validly notified within one month of that date;"),
    ("(c)", "provide a complete and compliant Article 15 response, comprising (i) a copy of all "
            "personal data undergoing processing that relates to our client, including extracts "
            "or intelligible summaries of relevant emails, Teams/Slack messages and logs rather "
            "than blanket exclusions, and (ii) all of the Article 15(1)(a)–(h) information "
            "(purposes, categories, recipients or categories of recipients, retention periods, "
            "source, the existence of any automated decision-making, and the safeguards for any "
            "transfers);"),
    ("(d)", "where any material is withheld or redacted, identify it and give specific reasons "
            "under Article 12(4), rather than relying on categorical exclusions;"),
    ("(e)", "carry out and document a personal-data-breach assessment under Articles 33 and 34 in "
            "respect of the member/union list and the BankID artifacts, and confirm whether "
            "notification to IMY and/or to affected data subjects has been or will be made;"),
    ("(f)", "explain the access controls, redaction process and lawful basis applied to the "
            "third-party personal data (including the personnummer, contact details, salaries, "
            "trade-union affiliation, BankID data and health data described above), and confirm "
            "the steps taken to secure or delete it; and"),
    ("(g)", "confirm the secure method and timetable by which the compliant response will be "
            "delivered to our client at his nominated personal address."),
]
for tag, text in reqs:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(1.0)
    para.paragraph_format.first_line_indent = Inches(-0.5)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(f"{tag}\t")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r2 = para.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(11)

# ============================ 6. CONSEQUENCES / RESERVATION ============================
doc.add_heading("6.  Consequences of non-compliance and reservation of rights", level=1)
numbered("6.1", "If you do not comply within the time stated, our client intends, without further "
         "notice, to (i) lodge a complaint with IMY under Article 77 GDPR; (ii) bring proceedings "
         "for an order requiring compliance with Article 15; and (iii) claim compensation under "
         "Article 82 GDPR for the material and non-material damage caused by the breaches "
         "described above, together with costs.")
numbered("6.2", "Our client expressly reserves all of his rights and remedies. Nothing in this "
         "letter, and no omission from it, shall be treated as a waiver of any right or as an "
         "exhaustive statement of his case. The breaches summarised above are without prejudice "
         "to further matters that may emerge once a compliant Article 15 response is provided.")
numbered("6.3", "We also remind you of your obligation not to destroy, alter or delete any data, "
         "logs, messages or documents relevant to this matter, including the exhibits referred to "
         "below. You should treat this letter as notice to preserve all such material.")

# ============================ 7. RESPONSE ============================
doc.add_heading("7.  Response", level=1)
numbered("7.1", "Please address your response to this firm, quoting our reference above. We look "
         "forward to hearing from you within 14 days.")

p("", space_after=6)
p("Yours faithfully,", space_after=24)
p("_____________________________", space_after=0)
p("[ Name ]", bold=True, space_after=0)
p("[ Title — e.g. Advokat / Partner ]", space_after=0)
p("for and on behalf of [ LAW FIRM NAME ]", space_after=0)
p("Counsel for Mr Thaer Saidi", italic=True, space_after=12)

# ============================ SCHEDULE ============================
doc.add_page_break()
doc.add_heading("Schedule — Exhibits relied upon", level=1)
p("The following items are relied upon in support of this letter and are available to be "
  "produced. Third-party identifiers have been redacted in our working copies.", space_after=8)

exhibits = [
    ("Exhibit A", "Spreadsheet listing 105 identified employees with personnummer, private "
                  "contact details, home addresses, salaries and trade-union club/section "
                  "(“member-list”)."),
    ("Exhibit B", "BankID signature/SDO records for our client and at least one third party, "
                  "containing personal identity numbers, device/IP identifiers and full "
                  "signature/certificate payloads (“BankIDSignatures”)."),
    ("Exhibit C", "Unit4 export of our client’s HR record: personnummer, bank account/IBAN, "
                  "employment terms, absence records and next-of-kin (spouse) details "
                  "(“TS Unit4 Data”)."),
    ("Exhibit D", "Extension notice dated 13 February 2026 issued by Nordcloud Hosting Sweden "
                  "AB, stating receipt on 27.1.2026 and a response deadline of 27.4.2026."),
    ("Exhibit E", "Internal Epical DSAR dossier showing the “Pending” data-source inventory, "
                  "the “mid-January” receipt date, handling by Mr Lars Rosenquist, and the "
                  "blanket scope exclusions."),
    ("Exhibit F", "Teams/email export disclosing a named third party’s heart attack and "
                  "sick-leave (health data)."),
    ("Exhibit G", "Teams export forwarding SFTP credentials, and an Azure role-assignment "
                  "export, evidencing insecure mixing of credentials with personal data."),
]
for tag, text in exhibits:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(1.3)
    para.paragraph_format.first_line_indent = Inches(-1.3)
    para.paragraph_format.space_after = Pt(6)
    r = para.add_run(f"{tag}:\t")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r2 = para.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(11)

p("", space_after=10)
hrule()
note = p("DRAFTING NOTE (delete before sending): Items in [square brackets] must be completed "
         "(firm details, addresses, file reference, signatory). Verify the true DSAR receipt "
         "date, the controller entity, and the exact contents/delivery of any disclosure already "
         "made before this letter is sent. This draft is prepared on the instructions and "
         "factual material provided and should be reviewed by the responsible advokat prior to "
         "service.", italic=True, size=9)

out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..",
                                   "Letter_Before_Action_Thaer_Saidi.docx"))
doc.save(out)
print("Saved:", out)
print("Paragraphs:", len(doc.paragraphs))
